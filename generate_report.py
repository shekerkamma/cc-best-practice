#!/usr/bin/env python3
"""Generate AI Strategy Report: Automotive Industry SAP OData Integration - Toyota TMNA Focus"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Styling helpers ──────────────────────────────────────────────────────────

FONT_NAME = "Calibri"
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
ACCENT_COLOR = RGBColor(0x2E, 0x75, 0xB6)   # SAP-ish blue
TABLE_HEADER_BG = "1B3A5C"
TABLE_ALT_BG = "E8EEF4"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = HEADING_COLOR
    return h


def add_body(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
        text_after = text
    else:
        text_after = text
    r2 = p.add_run(text_after)
    r2.font.name = FONT_NAME
    r2.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level == 1:
        p.style = doc.styles["List Bullet 2"]
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = FONT_NAME
    r2.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ── Document creation ────────────────────────────────────────────────────────

doc = Document()

# Default font
style = doc.styles["Normal"]
font = style.font
font.name = FONT_NAME
font.size = Pt(11)

# ============================================================================
# 1. COVER PAGE
# ============================================================================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AI Strategy:\nAutomotive Industry SAP OData Integration")
r.font.size = Pt(28)
r.font.name = FONT_NAME
r.font.color.rgb = HEADING_COLOR
r.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Toyota Motor North America (TMNA) Focus")
r.font.size = Pt(18)
r.font.name = FONT_NAME
r.font.color.rgb = ACCENT_COLOR

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run("May 2026")
r.font.size = Pt(14)
r.font.name = FONT_NAME
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
classification = doc.add_paragraph()
classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = classification.add_run("CONFIDENTIAL - STRATEGIC PLANNING")
r.font.size = Pt(10)
r.font.name = FONT_NAME
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.italic = True

doc.add_page_break()

# ============================================================================
# 2. TABLE OF CONTENTS
# ============================================================================

styled_heading(doc, "Table of Contents", level=1)
toc_items = [
    ("1.", "Executive Summary", 3),
    ("2.", "Market Signal Analysis", 4),
    ("3.", "Macro Thesis: Copilot vs Autopilot & Intelligence vs Judgement", 6),
    ("4.", "Market Sizing & Vertical Analysis", 8),
    ("5.", "Proof Points: Companies Winning in SAP + AI", 10),
    ("6.", "Operational Playbook: AI-Native SAP OData Integration", 12),
    ("7.", "Unit Economics: AI-Augmented vs Traditional SAP Operations", 14),
    ("8.", "Competitive Moats", 16),
    ("9.", "Risk Analysis", 17),
    ("10.", "Strategic Framework: Build vs Buy vs Partner", 19),
    ("11.", "Competitive Landscape: OEM AI + SAP Strategies", 21),
    ("12.", "References & Sources", 23),
]
for num, title_text, page in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{num} ")
    r1.bold = True
    r1.font.name = FONT_NAME
    r1.font.size = Pt(12)
    r2 = p.add_run(title_text)
    r2.font.name = FONT_NAME
    r2.font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# 3. EXECUTIVE SUMMARY
# ============================================================================

styled_heading(doc, "1. Executive Summary", level=1)

add_body(doc, "This report analyzes the strategic opportunity at the intersection of artificial intelligence, SAP enterprise systems, and OData API integration within the automotive industry, with a specific focus on Toyota Motor North America (TMNA). The analysis draws on market intelligence gathered in May 2026 across competitive dynamics, technology platforms, and operational economics.")

doc.add_paragraph()
styled_heading(doc, "Five Key Findings", level=2)

findings = [
    ("Finding 1 - Massive Market Convergence: ",
     "The automotive AI market is projected to reach $15.5B in 2024, growing at 15-37% CAGR depending on scope. SAP's installed base across automotive OEMs creates a unique integration surface for AI-native services, with OData APIs serving as the standardized bridge between legacy ERP systems and modern AI agents."),
    ("Finding 2 - Toyota TMNA Is Accelerating AI Adoption: ",
     "TMNA has committed $10B over five years to U.S. operations with AI as a strategic pillar, established a centralized Enterprise AI group, partnered with Deloitte and AWS for agentic AI in supply chain, and deployed Google Cloud-based AI tools that saved 10,000+ manufacturing man-hours annually."),
    ("Finding 3 - OData Is the De Facto AI Integration Layer for SAP: ",
     "SAP S/4HANA exposes 15,000+ OData API endpoints. The emergence of OData-to-MCP (Model Context Protocol) bridges enables AI agents to dynamically discover and interact with SAP business objects without custom coding, transforming the integration paradigm from middleware-heavy to API-native."),
    ("Finding 4 - Copilot-to-Autopilot Transition Creates a $340B+ Opportunity: ",
     "Applying Sequoia Capital's framework, automotive SAP operations are transitioning from copilot tools (AI-assisted analysts) to autopilot services (AI agents executing complete workflows). For every $1 spent on SAP software, $6 is spent on services -- all addressable by AI-native automation."),
    ("Finding 5 - First-Mover Advantage in Domain-Specific SAP AI Is Defensible: ",
     "Companies that build proprietary data moats from SAP operational data, combined with automotive domain expertise (IATF 16949, Toyota Production System, JIT/JIS logistics), will establish durable competitive advantages. The window for capturing this position is 18-24 months."),
]
for prefix, text in findings:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_page_break()

# ============================================================================
# 4. MARKET SIGNAL ANALYSIS
# ============================================================================

styled_heading(doc, "2. Market Signal Analysis", level=1)

styled_heading(doc, "2.1 SAP + AI Convergence", level=2)
add_body(doc, "The SAP ecosystem is undergoing a fundamental transformation as AI capabilities are embedded across the platform stack. Key signals include:")
add_bullet(doc, "SAP invested in Anthropic, Aleph Alpha, and Cohere as part of a $1B+ AI commitment through Sapphire Ventures, signaling commitment to multi-model AI integration.", bold_prefix="Strategic AI Investments: ")
add_bullet(doc, "SAP BTP AI Core provides enterprise-grade AI runtime, enabling custom model deployment alongside SAP business processes. Joule Studio and agent builder capabilities support multi-agent architectures at enterprise scale.", bold_prefix="SAP BTP AI Core: ")
add_bullet(doc, "SAP is transitioning from per-user subscription pricing to AI consumption-based pricing, reflecting the shift from human-operated to agent-operated enterprise systems.", bold_prefix="Pricing Model Shift: ")
add_bullet(doc, "SAP Integration Suite now includes Joule copilot-based discovery, integration flow generation, Groovy script optimization, and API anomaly detection.", bold_prefix="AI-Native Integration Suite: ")

styled_heading(doc, "2.2 Toyota TMNA Digital Initiatives", level=2)
add_body(doc, "Toyota Motor North America has made significant strategic moves in AI and digital transformation:")
add_bullet(doc, "TMNA partnered with Deloitte and AWS to deploy agentic AI directly into end-to-end supply chain workflows, targeting improvements in responsiveness, planning accuracy, and customer experience.", bold_prefix="Agentic AI for Supply Chain: ")
add_bullet(doc, "$670M Toyota Invention Partners fund for early-stage AI startups, plus $800M Woven Capital Fund II for growth-stage AI/automation companies, totaling $1.5B in new venture funding.", bold_prefix="$1.5B Venture Investment: ")
add_bullet(doc, "Established a centralized Enterprise AI group within TMNA, partnering with Toyota Connected North America (TCNA) to coordinate AI research across divisions, affiliates, and regions.", bold_prefix="Enterprise AI Organization: ")
add_bullet(doc, "In-house AI platform built on Google Cloud enables factory workers to create ML models without data science expertise. Model count grew from 8,000 (2023) to 10,000 (2024), saving 10,000+ man-hours annually.", bold_prefix="Factory AI Democratization: ")
add_bullet(doc, "Global AI Accelerator (GAIA) targets 11 key categories for AI adoption across the Toyota group, from ADAS and manufacturing to business software and customer relations.", bold_prefix="GAIA Initiative: ")

styled_heading(doc, "2.3 OData as the Integration Layer", level=2)
add_body(doc, "OData (Open Data Protocol) has emerged as the critical bridge between SAP systems and AI agents:")
add_bullet(doc, "OData abstracts SAP's internal complexity into RESTful endpoints with standardized query operations ($filter, $select, $expand), reducing payload size and latency while maintaining enterprise-grade security via OAuth 2.0.", bold_prefix="Standardized REST Access: ")
add_bullet(doc, "The OData MCP Proxy (open-source npm package) turns any OData V2 or REST API into an MCP server, automatically generating tools for every entity set and operation without custom code.", bold_prefix="OData-to-MCP Bridge: ")
add_bullet(doc, "SAP OData APIs connect to SAP systems' catalog, fetch available services, parse metadata, and generate necessary AI agent tools on the fly -- enabling dynamic discovery of 15,000+ endpoints.", bold_prefix="Dynamic Service Discovery: ")
add_bullet(doc, "SAP supports both OData V2 (majority of existing services) and V4 (newer services with improved performance, reduced payloads, and advanced metadata modeling).", bold_prefix="V2 and V4 Support: ")

doc.add_page_break()

# ============================================================================
# 5. MACRO THESIS
# ============================================================================

styled_heading(doc, "3. Macro Thesis: Strategic Frameworks Applied", level=1)

styled_heading(doc, "3.1 Sequoia Capital: Copilot vs Autopilot Framework", level=2)
add_body(doc, 'In March 2026, Sequoia Capital published "Services: The New Software," arguing that the next trillion-dollar company will sell the work itself rather than tools. This framework maps directly to automotive SAP integration:')

add_table(doc,
    ["Dimension", "Copilot (2024-2025)", "Autopilot (2026+)"],
    [
        ["Value Delivery", "AI assists SAP analysts with data queries, report generation, anomaly flagging", "AI agents autonomously execute procurement, logistics planning, quality workflows end-to-end"],
        ["Pricing Model", "Per-seat SaaS license for AI-enhanced SAP tools", "Per-outcome or per-transaction pricing tied to completed business processes"],
        ["Customer Relationship", "Customer still owns the work; AI accelerates human operators", "Customer buys completed outcomes; AI owns execution with human oversight"],
        ["SAP Integration", "Copilot overlays on SAP Fiori; read-heavy OData queries", "Full read-write OData agent loops; autonomous work order creation, PO processing"],
        ["Market Size", "Tool market: SAP add-on revenue ($2-5B)", "Services market: 6x software spend ($12-30B in automotive alone)"],
        ["Example", "Joule answering natural-language queries about inventory", "AI agent detecting supply disruption, renegotiating with suppliers, updating SAP MM/PP automatically"],
    ],
    col_widths=[3, 6, 6]
)

add_body(doc, "The automotive industry is currently in the copilot phase, with solutions like SAP Joule and third-party AI assistants augmenting human SAP operators. The transition to autopilot -- where AI agents execute complete business processes via OData APIs -- represents a 6x expansion in addressable market.")

styled_heading(doc, "3.2 Emergence Capital: Intelligence vs Judgement Framework", level=2)
add_body(doc, "Emergence Capital's AI-Native Services Playbook distinguishes between tasks requiring intelligence (pattern recognition, data processing) and those requiring judgement (contextual decisions with incomplete information). For automotive SAP operations:")

add_table(doc,
    ["Task Category", "Intelligence (AI-Native)", "Judgement (Human-in-Loop)"],
    [
        ["Procurement", "Automated PO matching, invoice verification, spend analysis via SAP MM OData", "Supplier selection for critical components, strategic sourcing decisions"],
        ["Manufacturing", "Predictive maintenance scheduling, quality defect detection, OEE optimization", "Production line rebalancing during model changeovers, safety protocol updates"],
        ["Supply Chain", "Demand forecasting, inventory optimization, logistics routing via SAP APO", "Supplier risk assessment during geopolitical disruptions, sole-source decisions"],
        ["Quality", "SPC data analysis, IATF 16949 compliance monitoring, CAPA tracking", "Root cause determination for systemic quality issues, customer notification decisions"],
        ["Finance", "Automated journal entries, intercompany reconciliation, variance analysis", "Capital allocation for new EV production lines, M&A due diligence"],
    ],
    col_widths=[2.5, 5.5, 5.5]
)

add_body(doc, "The key insight: AI-native services should focus on intelligence tasks first (1-2 jobs-to-be-done), building domain credibility before expanding into judgement-adjacent areas. For SAP OData integration, this means starting with high-volume, data-rich processes like procurement matching and predictive maintenance before tackling strategic supply chain decisions.")

doc.add_page_break()

# ============================================================================
# 6. MARKET SIZING & VERTICAL ANALYSIS
# ============================================================================

styled_heading(doc, "4. Market Sizing & Vertical Analysis", level=1)

styled_heading(doc, "4.1 Automotive AI Total Addressable Market", level=2)
add_body(doc, "Market research firms project the automotive AI market with varying scope definitions:")

add_table(doc,
    ["Source", "2024 Estimate", "2030 Projection", "CAGR"],
    [
        ["Markets and Markets", "$15.5B", "$38.5B", "15.3%"],
        ["Grand View Research", "$4.3B", "$14.9B", "23.4%"],
        ["GM Insights", "$4.8B", "$18.2B", "25.1%"],
        ["Market.us", "$7.7B", "$134.3B (2033)", "37.4%"],
        ["Precedence Research", "$4.7B (2025)", "$48.6B (2034)", "26.2%"],
    ],
    col_widths=[4, 3, 3, 2.5]
)

add_body(doc, "The variance reflects different scope definitions: narrower estimates focus on embedded AI in vehicles (ADAS, infotainment), while broader estimates include manufacturing AI, supply chain optimization, and enterprise operations. For this analysis, the enterprise/operations segment is most relevant.")

styled_heading(doc, "4.2 SAP Integration TAM", level=2)
add_body(doc, "The SAP ecosystem supports a massive integration market:")
add_bullet(doc, "SAP serves 77% of the world's transaction revenue, with 400,000+ customers globally. In automotive, virtually every major OEM and Tier 1 supplier runs SAP.", bold_prefix="SAP Market Penetration: ")
add_bullet(doc, "SAP Integration Suite, middleware platforms, and custom integration development represent a $15-20B annual market, with automotive accounting for approximately 12-15% ($1.8-3.0B).", bold_prefix="Integration Spend: ")
add_bullet(doc, "With SAP's shift to consumption-based AI pricing and the emergence of AI agent platforms, integration spend is transitioning from capital expenditure (custom ABAP development, middleware licenses) to operational expenditure (per-transaction AI agent fees).", bold_prefix="CapEx to OpEx Shift: ")

styled_heading(doc, "4.3 OData-Specific Opportunity", level=2)
add_body(doc, "The OData API layer creates a distinct, defensible market segment:")
add_bullet(doc, "SAP S/4HANA exposes 15,000+ OData endpoints, each representing a potential AI agent interaction point. At an average of $0.01-0.10 per AI-mediated transaction, even 1% penetration of SAP transaction volume represents a $500M-2B opportunity.", bold_prefix="API Surface Area: ")
add_bullet(doc, "The OData-to-MCP bridge pattern eliminates the need for custom connector development, reducing integration time from months to days and cost from $500K+ to under $50K per connection.", bold_prefix="Integration Acceleration: ")
add_bullet(doc, "Automotive-specific OData services (vehicle management, parts catalogs, dealer networks, warranty management) require domain expertise that creates natural barriers to entry.", bold_prefix="Domain Specificity: ")

doc.add_page_break()

# ============================================================================
# 7. PROOF POINTS
# ============================================================================

styled_heading(doc, "5. Proof Points: Companies Winning in SAP + AI", level=1)

styled_heading(doc, "5.1 Celonis (Process Mining + AI)", level=2)
add_bullet(doc, "$13B peak valuation, $1.77B total funding raised", bold_prefix="Valuation: ")
add_bullet(doc, "Sits on top of SAP, Oracle, and Salesforce to mine process execution data and identify optimization opportunities", bold_prefix="Strategy: ")
add_bullet(doc, "Ranked #3 on Fortune Future 50 (2025); IPO expected 2026", bold_prefix="Traction: ")
add_bullet(doc, "Demonstrates that the data exhaust from SAP operations is extraordinarily valuable -- process mining is the 'intelligence' layer that feeds AI automation", bold_prefix="Relevance to Thesis: ")

styled_heading(doc, "5.2 SAP LeanIX (IT Architecture Intelligence)", level=2)
add_bullet(doc, "Acquired by SAP in November 2023 after raising $128M independently", bold_prefix="Acquisition: ")
add_bullet(doc, "Enterprise architecture management that maps IT landscapes, enabling informed decisions about application rationalization and cloud migration", bold_prefix="Strategy: ")
add_bullet(doc, "SAP acquiring LeanIX validates the strategic value of metadata and architecture intelligence as a precursor to AI-driven enterprise operations", bold_prefix="Relevance to Thesis: ")

styled_heading(doc, "5.3 Automation Anywhere", level=2)
add_bullet(doc, "Advanced analytics with IQ Bot for unstructured data processing, deep integrations with SAP, Oracle, and Salesforce", bold_prefix="Capabilities: ")
add_bullet(doc, "Multi-system workflow automation that scales across departments and integrates with legacy SAP systems via RFC/BAPI and OData connectors", bold_prefix="SAP Integration: ")
add_bullet(doc, "Proves that RPA-to-AI evolution in SAP environments is viable and scalable, though limited by bot-level (not agent-level) intelligence", bold_prefix="Relevance to Thesis: ")

styled_heading(doc, "5.4 Workato (Enterprise iPaaS + AI Agents)", level=2)
add_bullet(doc, "Connects to 1,200+ enterprise applications including deep SAP integration via RFC and OData", bold_prefix="Platform: ")
add_bullet(doc, 'Launched "Genies" -- AI agents that independently retrieve information, make decisions, and perform actions across applications. Now positioned as "Enterprise MCP for Agentic AI."', bold_prefix="AI Evolution: ")
add_bullet(doc, "Validates the platform play: combining integration infrastructure with AI agent capabilities for SAP-connected workflows", bold_prefix="Relevance to Thesis: ")

styled_heading(doc, "5.5 Automaker Internal Initiatives", level=2)
add_bullet(doc, "Partnered with Deloitte and AWS for agentic AI in supply chain; centralized Enterprise AI group; $10B five-year U.S. investment with AI focus; Google Cloud-based factory AI platform", bold_prefix="Toyota TMNA: ")
add_bullet(doc, "$1.1B AI initiative; 'WE & AI' training for 130,000+ employees; Catena-X data sharing platform with SAP; Large Industry Models (LIM) for manufacturing; ChatGPT integration in vehicles", bold_prefix="Volkswagen: ")
add_bullet(doc, "AI in product development saving weeks of engineering effort; manufacturing AI for quality and efficiency", bold_prefix="Ford: ")
add_bullet(doc, "Catena-X founding member with SAP; exploring cross-manufacturer AI data sharing for supply chain optimization", bold_prefix="BMW: ")

doc.add_page_break()

# ============================================================================
# 8. OPERATIONAL PLAYBOOK
# ============================================================================

styled_heading(doc, "6. Operational Playbook: AI-Native SAP OData Integration", level=1)

styled_heading(doc, "6.1 Reference Architecture", level=2)
add_body(doc, "The target architecture for AI-native SAP OData integration consists of four layers:")

add_table(doc,
    ["Layer", "Components", "Purpose"],
    [
        ["1. SAP Core", "SAP S/4HANA, SAP ECC, SAP BW/4HANA", "System of record for business transactions, master data, and analytics"],
        ["2. API Gateway", "SAP Gateway, OData V2/V4 services, OAuth 2.0, SAP API Business Hub", "Standardized REST access to SAP business objects with security and rate limiting"],
        ["3. AI Agent Platform", "MCP servers, OData-MCP Proxy, Agent orchestration, LLM inference (Claude/GPT)", "Dynamic discovery of SAP services, tool generation, agent reasoning and execution"],
        ["4. Domain Intelligence", "Automotive-specific models, IATF 16949 rules engine, TPS knowledge base", "Domain expertise layer that guides agent decisions with automotive context"],
    ],
    col_widths=[2.5, 5, 5.5]
)

styled_heading(doc, "6.2 Data Flow Patterns", level=2)
add_body(doc, "Three primary data flow patterns for AI-SAP integration via OData:")

add_bullet(doc, "AI agent queries SAP via OData $filter and $select to retrieve specific business objects (purchase orders, production orders, quality records). Used for: dashboarding, anomaly detection, compliance monitoring. Latency target: <500ms per query.", bold_prefix="Pattern 1 -- Read-Query-Analyze: ")
add_bullet(doc, "AI agent reads SAP data, applies ML models or LLM reasoning, then writes back to SAP via OData POST/PATCH operations. Used for: predictive maintenance work orders, automated PO creation, quality hold decisions. Requires: transactional integrity, audit logging, rollback capability.", bold_prefix="Pattern 2 -- Read-Reason-Write: ")
add_bullet(doc, "Multiple AI agents coordinate across SAP modules (MM, PP, QM, FI) to execute end-to-end business processes. Used for: supplier disruption response (detect in MM, replan in PP, adjust in FI), new model launch logistics. Requires: agent orchestration framework, conflict resolution, human escalation paths.", bold_prefix="Pattern 3 -- Multi-Agent Orchestration: ")

styled_heading(doc, "6.3 Agent Design Patterns", level=2)
add_bullet(doc, "Single-purpose agents with deep expertise in one SAP module (e.g., MM Procurement Agent, PP Production Planning Agent, QM Quality Agent). Each agent has its own OData service catalog and domain-specific reasoning.", bold_prefix="Specialist Agents: ")
add_bullet(doc, "Coordinates specialist agents for cross-functional workflows. Maintains process state, handles agent-to-agent communication, and manages human escalation. Implements the Sequoia 'autopilot' pattern.", bold_prefix="Orchestrator Agent: ")
add_bullet(doc, "Lightweight agents that monitor OData event streams for specific conditions (threshold breaches, SLA violations, quality alerts) and trigger specialist agents when intervention is needed.", bold_prefix="Sentinel Agents: ")

doc.add_page_break()

# ============================================================================
# 9. UNIT ECONOMICS
# ============================================================================

styled_heading(doc, "7. Unit Economics: AI-Augmented vs Traditional SAP Operations", level=1)

styled_heading(doc, "7.1 Cost Comparison", level=2)

add_table(doc,
    ["Cost Dimension", "Traditional (ABAP/Middleware)", "AI-Native (OData + Agents)", "Delta"],
    [
        ["Integration Development", "$150-500K per connector (3-6 months)", "$10-50K per connector (1-2 weeks)", "70-90% reduction"],
        ["Maintenance Cost/Year", "$50-150K per integration (ABAP developer)", "$5-15K per agent (monitoring + model updates)", "80-90% reduction"],
        ["Transaction Processing", "$0.50-2.00 per human-processed transaction", "$0.01-0.10 per AI-processed transaction", "90-98% reduction"],
        ["Error Rate", "2-5% (manual data entry, process deviations)", "0.1-0.5% (AI validation with human escalation)", "90% reduction"],
        ["Time to Insight", "Hours to days (report requests, analyst queues)", "Seconds to minutes (real-time OData queries)", "95%+ reduction"],
        ["Scalability", "Linear (more transactions = more staff)", "Near-zero marginal cost per additional transaction", "Exponential leverage"],
    ],
    col_widths=[3, 4, 4, 2.5]
)

styled_heading(doc, "7.2 Development Velocity", level=2)
add_body(doc, "AI-native SAP integration fundamentally changes development economics:")
add_bullet(doc, "Traditional ABAP custom development requires 6-12 months for complex integrations. OData-based AI agent development reduces this to 2-4 weeks, with the OData-MCP bridge eliminating boilerplate connector code entirely.", bold_prefix="Speed: ")
add_bullet(doc, "Traditional: Senior ABAP consultant ($200-350/hr) x 1,000+ hours = $200K-350K. AI-native: ML engineer ($150-250/hr) x 80-160 hours = $12K-40K, plus LLM inference costs of $500-2,000/month.", bold_prefix="Cost per Integration: ")
add_bullet(doc, "OData APIs remain stable across SAP releases, reducing regression testing burden. AI agents self-adapt to schema changes via dynamic metadata discovery.", bold_prefix="Maintenance: ")

styled_heading(doc, "7.3 ROI Benchmarks from Manufacturing AI", level=2)
add_bullet(doc, "30-50% reduction in unplanned downtime; 20-40% extension of equipment life; maintenance cost savings of 20-25%. One automotive manufacturer saved $3.2M annually across 200+ CNC machines.", bold_prefix="Predictive Maintenance: ")
add_bullet(doc, "Scrap rate reductions of 20-35%; first-pass quality yields improved to 98.7%", bold_prefix="Quality Inspection: ")
add_bullet(doc, "300-500% ROI typical; breakeven within 3-6 months from a single prevented major failure; full program payback in 6-14 months", bold_prefix="Overall Manufacturing AI ROI: ")
add_bullet(doc, "AI cobots generating $4.8M annual value per assembly line with 195-day payback period", bold_prefix="Collaborative Robotics: ")

doc.add_page_break()

# ============================================================================
# 10. COMPETITIVE MOATS
# ============================================================================

styled_heading(doc, "8. Competitive Moats", level=1)

styled_heading(doc, "8.1 Data Moat from SAP Operational Data", level=2)
add_body(doc, "SAP systems contain the most valuable operational data in any enterprise: actual transaction records, supplier performance history, production quality metrics, and financial outcomes. An AI system that processes this data builds compound advantages:")
add_bullet(doc, "Every transaction processed by an AI agent improves its domain model. A procurement agent that has processed 1M+ POs across 50+ Toyota suppliers develops supplier-specific intelligence that cannot be replicated without equivalent data access.", bold_prefix="Learning Loop: ")
add_bullet(doc, "Celonis proved that mining SAP process data creates $13B+ in enterprise value. AI agents that both mine and act on this data capture value at both the intelligence and execution layers.", bold_prefix="Process Intelligence: ")
add_bullet(doc, "As AI agents build history within a customer's SAP environment, switching costs increase exponentially. The agent's learned context (supplier behaviors, seasonal patterns, quality correlations) is non-transferable.", bold_prefix="Switching Costs: ")

styled_heading(doc, "8.2 OData API Expertise", level=2)
add_bullet(doc, "Deep understanding of SAP's 15,000+ OData endpoints, their behaviors, performance characteristics, and edge cases is rare expertise that takes years to develop.", bold_prefix="Technical Depth: ")
add_bullet(doc, "Automotive SAP configurations (vehicle variant management, JIT/JIS delivery, warranty claims processing) require specialized OData service knowledge.", bold_prefix="Industry-Specific API Knowledge: ")
add_bullet(doc, "SAP OData APIs have undocumented behaviors, performance cliffs, and version-specific quirks. Accumulated operational knowledge of these creates a practical moat.", bold_prefix="Undocumented Behaviors: ")

styled_heading(doc, "8.3 Domain-Specific AI Models", level=2)
add_bullet(doc, "AI models trained on automotive manufacturing data (vibration patterns, torque curves, paint quality metrics) combined with SAP master data create cross-domain intelligence unavailable to general-purpose AI platforms.", bold_prefix="Manufacturing Domain Models: ")
add_bullet(doc, "Toyota Production System principles (Jidoka, Kaizen, Heijunka) encoded into AI agent decision-making create culturally aligned automation that generic agents cannot match.", bold_prefix="TPS-Aligned Agent Behavior: ")
add_bullet(doc, "IATF 16949 compliance rules, PPAP requirements, and automotive-specific quality methodologies (8D, FMEA) embedded in agent logic create regulatory defensibility.", bold_prefix="Regulatory Knowledge: ")

doc.add_page_break()

# ============================================================================
# 11. RISK ANALYSIS
# ============================================================================

styled_heading(doc, "9. Risk Analysis", level=1)

styled_heading(doc, "9.1 Mirage PMF Risks", level=2)
add_body(doc, "The risk of 'mirage product-market fit' -- where early adoption signals mask fundamental scalability issues -- is significant in enterprise AI:")
add_bullet(doc, "POC projects may show impressive demos on clean data subsets but fail when exposed to the full complexity of production SAP environments (custom fields, legacy data, edge-case transactions).", bold_prefix="POC-to-Production Gap: ")
add_bullet(doc, "Executive sponsors may champion AI initiatives based on vendor marketing rather than operational reality. When AI agents make errors in production SAP systems, organizational trust can collapse rapidly.", bold_prefix="Executive Sponsorship Risk: ")
add_bullet(doc, "AI may automate the wrong processes -- those that are visible and demonstrable rather than those with the highest operational impact.", bold_prefix="Value Attribution: ")

styled_heading(doc, "9.2 SAP Vendor Lock-in", level=2)
add_bullet(doc, "SAP's shift to consumption-based AI pricing and its May 2026 API policy changes force enterprises to reckon with data ownership in the age of agentic AI.", bold_prefix="API Policy Changes: ")
add_bullet(doc, "Building AI agents deeply integrated with SAP OData creates dependency on SAP's API stability, versioning decisions, and pricing strategy.", bold_prefix="Platform Dependency: ")
add_bullet(doc, "SAP's own AI capabilities (Joule, AI Core, agent builder) may compete directly with third-party solutions, with SAP potentially restricting API access to favor its own offerings.", bold_prefix="Competitive Displacement: ")

styled_heading(doc, "9.3 Regulatory Risks (IATF 16949)", level=2)
add_bullet(doc, "IATF 16949:2016 requires documented quality management systems with full traceability. AI agent decisions in quality-critical processes must be auditable, explainable, and traceable to certified procedures.", bold_prefix="Quality Traceability: ")
add_bullet(doc, "Future IATF revisions may incorporate specific requirements for AI/ML systems in automotive quality management, potentially requiring certification of AI models used in production decisions.", bold_prefix="AI-Specific Regulations: ")
add_bullet(doc, "Automotive OEMs face varying data protection requirements across jurisdictions (GDPR, CCPA, China's PIPL). AI agents processing SAP data across Toyota's global operations must comply with all applicable data sovereignty requirements.", bold_prefix="Cross-Border Compliance: ")

styled_heading(doc, "9.4 Data Sovereignty", level=2)
add_bullet(doc, "AI agents processing SAP data must respect data residency requirements. Toyota's global operations span jurisdictions with conflicting data localization laws.", bold_prefix="Data Residency: ")
add_bullet(doc, "Sending SAP operational data to cloud-based LLMs raises concerns about intellectual property exposure, particularly for manufacturing process data that embodies Toyota Production System know-how.", bold_prefix="IP Exposure: ")
add_bullet(doc, "Depending on LLM training data handling, SAP operational data processed by AI agents could theoretically influence model training, leaking competitive intelligence.", bold_prefix="Model Training Data Risks: ")

doc.add_page_break()

# ============================================================================
# 12. STRATEGIC FRAMEWORK
# ============================================================================

styled_heading(doc, "10. Strategic Framework: Build vs Buy vs Partner", level=1)

add_body(doc, "For Toyota TMNA, the build-vs-buy-vs-partner decision for AI-native SAP OData integration depends on strategic importance, time-to-value, and long-term competitive differentiation.")

styled_heading(doc, "10.1 Decision Matrix", level=2)

add_table(doc,
    ["Dimension", "Build (Internal)", "Buy (Vendor)", "Partner (Joint Venture)"],
    [
        ["Time to Market", "18-36 months", "3-6 months", "6-12 months"],
        ["Cost (Year 1)", "$5-15M (team + infrastructure)", "$1-3M (licenses + implementation)", "$3-8M (shared investment)"],
        ["IP Ownership", "Full ownership", "Vendor retains IP; TMNA licenses", "Shared IP with negotiated terms"],
        ["Customization", "Unlimited; tailored to TPS", "Limited to vendor platform", "Moderate; influenced by both parties"],
        ["Data Control", "Complete control", "Data processing by vendor", "Governed by partnership agreement"],
        ["Scalability Risk", "Talent acquisition bottleneck", "Vendor dependency", "Partnership alignment risk"],
        ["Competitive Moat", "Strong (proprietary TPS + AI)", "Weak (competitors buy same tool)", "Moderate (shared advantage)"],
        ["Best For", "Core differentiating capabilities", "Non-differentiating, commodity AI", "Strategic capabilities with shared risk"],
    ],
    col_widths=[2.5, 4, 4, 4]
)

styled_heading(doc, "10.2 Recommended Approach: Hybrid Strategy", level=2)
add_body(doc, "A three-tier approach is recommended for TMNA:")

add_bullet(doc, "Build internally for AI capabilities that directly embody Toyota Production System principles and create lasting competitive advantage. This includes: TPS-aligned manufacturing AI agents, proprietary quality prediction models trained on Toyota-specific data, and supply chain intelligence that leverages Toyota's unique supplier relationships.", bold_prefix="Tier 1 -- Build (Core Differentiators): ")
add_bullet(doc, "Buy best-of-breed solutions for commodity AI capabilities that don't differentiate. This includes: general-purpose SAP integration platforms (Workato, Celonis), standard predictive maintenance modules (DXC SPARK AI), and financial automation (invoice processing, reconciliation).", bold_prefix="Tier 2 -- Buy (Commodity Capabilities): ")
add_bullet(doc, "Partner for capabilities that require scale beyond TMNA's resources but are strategically important. This includes: Catena-X/Cofinity-X participation for cross-manufacturer data sharing, Deloitte/AWS engagement for agentic AI at enterprise scale, and potential joint ventures with AI-native services companies for automotive-specific SAP agent platforms.", bold_prefix="Tier 3 -- Partner (Strategic Scale): ")

doc.add_page_break()

# ============================================================================
# 13. COMPETITIVE LANDSCAPE
# ============================================================================

styled_heading(doc, "11. Competitive Landscape: OEM AI + SAP Strategies", level=1)

add_table(doc,
    ["OEM", "AI Strategy", "SAP Integration", "Key Differentiator", "Risk Factor"],
    [
        ["Toyota TMNA", "Centralized Enterprise AI group; $10B U.S. investment; Deloitte/AWS agentic AI partnership; GAIA global accelerator", "SAP S/4HANA across global operations; IBM-led modernization; OData-based integration roadmap", "TPS + AI fusion; factory worker AI democratization (10K+ ML models); $1.5B venture funds", "Organizational complexity across affiliates; conservative culture may slow agent adoption"],
        ["Volkswagen Group", "$1.1B AI initiative; 130K employees trained; Large Industry Models (LIM); ChatGPT in vehicles", "SAP founding partner in Catena-X; cross-manufacturer data sharing; enterprise SAP landscape", "Scale of AI training program; Catena-X leadership; willingness to invest heavily", "Software subsidiary CARIAD challenges; organizational restructuring distractions"],
        ["BMW", "Catena-X co-founder; AI in manufacturing quality and logistics", "Deep SAP integration; Catena-X data sharing with SAP as platform partner", "Premium brand operational excellence; strong IT organization", "Smaller scale limits AI investment; premium segment margin pressure from EVs"],
        ["Ford", "AI in product development saving weeks of effort; manufacturing AI focus", "SAP operations across enterprise; modernization in progress", "Engineering AI for vehicle design; agile organizational transformation", "Financial pressure; EV transition costs limiting AI investment bandwidth"],
        ["GM", "AI across manufacturing and customer experience", "Enterprise SAP footprint; integration modernization", "Cruise autonomous driving AI expertise applicable to enterprise", "Cruise setbacks may have dampened AI ambition; leadership transitions"],
        ["Hyundai-Kia", "Rapid AI adoption across manufacturing and connected services", "SAP operations with increasing cloud adoption", "Speed of execution; vertical integration advantage", "Less mature enterprise AI organization compared to Toyota/VW"],
    ],
    col_widths=[2, 4, 3.5, 3.5, 3]
)

styled_heading(doc, "11.1 Competitive Analysis Summary", level=2)
add_body(doc, "Toyota TMNA is positioned as a leader in automotive enterprise AI, distinguished by three factors:")
add_bullet(doc, "No other OEM has matched Toyota's approach of enabling factory workers to create their own ML models. With 10,000+ models created in 2024, this bottom-up AI adoption creates organic demand for SAP integration as workers seek to connect their models with enterprise data.", bold_prefix="1. Democratized AI Adoption: ")
add_bullet(doc, "Toyota's 70+ year history with lean manufacturing provides a unique framework for AI agent design. AI agents that embody Jidoka (stopping to fix problems), Kaizen (continuous improvement), and Heijunka (production leveling) will be fundamentally different from generic enterprise AI.", bold_prefix="2. TPS as AI Design Philosophy: ")
add_bullet(doc, "With $1.5B in venture funding specifically targeting AI/automation startups, plus the $10B U.S. investment, TMNA has more financial firepower directed at AI than any competitor except Volkswagen.", bold_prefix="3. Investment Scale: ")

doc.add_page_break()

# ============================================================================
# 14. REFERENCES & SOURCES
# ============================================================================

styled_heading(doc, "12. References & Sources", level=1)

references = {
    "Toyota & TMNA Digital Transformation": [
        ("CIMData", "Toyota Motor Corporation Chooses SAP S/4HANA and SAP HANA", "https://www.cimdata.com/de/industry-summary-articles/item/10519-toyota-motor-corporation-chooses-sap-s-4hana-and-sap-hana-to-help-drive-operational-efficiency"),
        ("Process Excellence Network", "Toyota Accelerates Digital Transformation", "https://www.processexcellencenetwork.com/digital-transformation/news/toyota-accelerates-digital-transformation-with-sap-s4hana-upgrades"),
        ("SAP Africa News Center", "IBM Accelerates Digital Transformation for Toyota South Africa Motors", "https://news.sap.com/africa/2025/11/ibm-accelerates-digital-transformation-for-toyota-south-africa-motors-with-sap-s-4hana-upgrades/"),
        ("ERP Today", "IBM Pushes Toyota South Africa's SAP Modernization Forward", "https://erp.today/ibm-pushes-toyota-south-africas-sap-modernization-forward-with-on-time-s-4hana-upgrades/"),
        ("SiliconANGLE", "Agentic AI at Scale: AWS, Deloitte and Toyota TMNA", "https://siliconangle.com/2025/12/19/aws-deloitte-toyota-motor-north-americas-shift-agentic-ai-awsreinvent/"),
        ("Google Cloud Blog", "How Toyota is Revolutionizing Manufacturing with AI", "https://cloud.google.com/blog/topics/hybrid-cloud/toyota-ai-platform-manufacturing-efficiency"),
        ("Chief AI Officer Blog", "How Toyota Gave AI Tools to Factory Workers and Saved 10,000 Hours", "https://chiefaiofficer.com/blog/how-toyota-gave-ai-tools-to-factory-workers-and-saved-10000-hours/"),
        ("Enki AI", "Toyota's AI Strategy 2026: The Definitive Deep Dive", "https://enkiai.com/ai-market-intelligence/toyotas-ai-strategy-2026-the-definitive-deep-dive/"),
        ("Toyota Newsroom", "Toyota and Generative AI: It's Here, and This is How We're Using It", "https://pressroom.toyota.com/toyota-and-generative-ai-its-here-and-this-is-how-were-using-it/"),
        ("Toyota Global Newsroom", "Five Toyota Group Companies to Accelerate AI and Software", "https://global.toyota/en/newsroom/corporate/42805724.html"),
    ],
    "SAP OData & AI Integration": [
        ("SAP Help", "OData API - SAP Integration Suite", "https://help.sap.com/docs/integration-suite/sap-integration-suite/odata-api"),
        ("Medium (Alice Vinogradova)", "I Built the Universal OData-MCP Bridge", "https://medium.com/@elfee/i-built-the-universal-odata-mcp-bridge-and-it-actually-works-20c9e34c9a87"),
        ("DEV Community", "4 Trends Redefining SAP-Salesforce Integration: AI, Cloud, OData, Low-Code", "https://dev.to/builderyo/the-4-trends-redefining-sap-salesforce-integration-ai-cloud-odata-and-low-code-f00"),
        ("SAP Community", "Integration Suite in 2025 & 2026", "https://community.sap.com/t5/technology-blog-posts-by-members/integration-suite-in-2025-amp-2026/ba-p/14320912"),
        ("SAP Community", "OData MCP Proxy - Introduction", "https://community.sap.com/t5/technology-blog-posts-by-members/odata-mcp-proxy-introduction/ba-p/14348684"),
        ("GitHub (Azure Samples)", "SAP OData API Guide Copilot", "https://github.com/Azure-Samples/sap-odata-api-guide-copilot"),
        ("SkyWork AI", "Unlocking SAP Data for AI: OData MCP Server Deep Dive", "https://skywork.ai/skypage/en/sap-data-ai-deep-dive/1980511314107813888"),
        ("Kai Waehner", "Data Ownership in the Age of Agentic AI: SAP's API Policy", "https://www.kai-waehner.de/blog/2026/05/02/data-ownership-in-the-age-of-agentic-ai-why-saps-api-policy-forces-a-data-integration-reckoning-for-every-enterprise/"),
    ],
    "Automotive AI Market": [
        ("Markets and Markets", "Automotive AI Market Report 2025-2030", "https://www.marketsandmarkets.com/Market-Reports/automotive-artificial-intelligence-market-248804391.html"),
        ("Grand View Research", "Automotive AI Market Industry Report 2030", "https://www.grandviewresearch.com/industry-analysis/automotive-artificial-intelligence-market-report"),
        ("GM Insights", "AI in Automotive Market Size & Share 2025-2034", "https://www.gminsights.com/industry-analysis/artificial-intelligence-ai-in-automotive-market"),
        ("Precedence Research", "Automotive AI Market Size to Hit $48.59B by 2034", "https://www.precedenceresearch.com/automotive-artificial-intelligence-market"),
    ],
    "SAP BTP & AI Core": [
        ("SAP News", "Customer-Specific AI Will Define the Next Era of Automotive", "https://news.sap.com/2026/02/customer-specific-ai-next-era-automotive-ecosystem/"),
        ("SAP Learning", "Maximizing Value for Automotive with SAP Business AI", "https://learning.sap.com/courses/introducing-the-automotive-industry/maximizing-value-for-the-automotive-industry-with-sap-business-ai"),
        ("SAP Discovery Center", "SAP AI Core Service Catalog", "https://discovery-center.cloud.sap/serviceCatalog/sap-ai-core"),
        ("SAVIC Technologies", "SAP AI Core - Enterprise AI Runtime on BTP", "https://www.savictech.com/data-ai/sap-ai-core/"),
    ],
    "Enterprise AI Partnerships": [
        ("TechCrunch", "Anthropic and OpenAI Launch Joint Ventures for Enterprise AI", "https://techcrunch.com/2026/05/04/anthropic-and-openai-are-both-launching-joint-ventures-for-enterprise-ai-services/"),
        ("SAP News", "SAP Advances Business AI with Investments in Anthropic, Aleph Alpha, Cohere", "https://news.sap.com/2023/07/generative-ai-investments-aleph-alpha-anthropic-cohere/"),
        ("Axios", "OpenAI and Anthropic Partner with Private Equity", "https://www.axios.com/2026/05/04/openai-anthropic-private-equity-enterprise-business"),
        ("PYMNTS", "Anthropic Launches Enterprise AI Firm With Wall Street Giants", "https://www.pymnts.com/artificial-intelligence-2/2026/anthropic-launches-enterprise-ai-firm-with-wall-street-giants/"),
    ],
    "Competitive Landscape": [
        ("Volkswagen Group", "Boosting Innovation: Volkswagen Group Invests in AI", "https://www.volkswagen-group.com/en/press-releases/boosting-innovation-reshaping-mobility-volkswagen-group-invests-in-ai-19852"),
        ("CIO Dive", "Volkswagen Adds 5 Years to AWS Cloud, AI Partnership", "https://www.ciodive.com/news/volkswagen-group-AWS-partnership-AI-cost-savings/758892/"),
        ("ADT Media", "Volkswagen's AI Strategy: Revolutionising Efficiency", "https://www.adt.media/software-defined-vehicles/volkswagen-wants-to-democratise-ainbsp-and-save-billions/648788"),
        ("Celonis", "Celonis Raises $290M for AI-Powered Process Mining", "https://www.celonis.com/news/press/venture-beat-celonis-raises-290-million-for-ai-powered-process-mining-at-2-5-billion-valuation"),
        ("Contrary Research", "Celonis Business Breakdown & Founding Story", "https://research.contrary.com/company/celonis"),
    ],
    "Strategic Frameworks": [
        ("Sequoia Capital", "Services: The New Software", "https://sequoiacap.com/article/services-the-new-software/"),
        ("Emergence Capital", "The AI-Native Services Playbook", "https://www.emcap.com/thoughts/the-ai-enabled-services-playbook"),
        ("Emergence Capital", "AI-Native Services: The Definitive Guide", "https://www.emcap.com/ai-native-services"),
        ("Emergence Capital", "Why AI-Native Services, and Why Now", "https://www.emcap.com/thoughts/why-ai-native-services-and-why-now"),
    ],
    "Manufacturing AI & ROI": [
        ("OxMaint", "ROI of AI Predictive Maintenance in Manufacturing", "https://www.oxmaint.com/blog/post/roi-ai-predictive-maintenance-manufacturing-cost-savings-analysis"),
        ("Markovate", "ROI of AI in Manufacturing", "https://markovate.com/roi-of-ai-in-manufacturing/"),
        ("Bain & Company", "Technology Is Radically Reshaping Auto Economics", "https://www.bain.com/insights/technology-is-radically-reshaping-auto-economics/"),
        ("Automotive Manufacturing Solutions", "How AI and Robotics Are Reshaping the Automotive Factory Floor", "https://www.automotivemanufacturingsolutions.com/automation/how-ai-and-nextgeneration-robotics-are-reshaping-the-automotive-factory-floor/2651815"),
        ("Automotive Manufacturing Solutions", "Toyota's 2025 Lean Manufacturing Evolution", "https://www.automotivemanufacturingsolutions.com/editors-pick/inside-toyotas-2025-leaner-manufacturing-system/645325"),
        ("Acropolium", "AI Agent Unit Economics: TCO, ROI, Payback", "https://acropolium.com/blog/ai-agent-unit-economics/"),
    ],
    "SAP Integration & Middleware": [
        ("Bizdata360", "SAP Integration Suite: Ultimate Guide 2026", "https://www.bizdata360.com/sap-integration-suite/"),
        ("SAP", "Enterprise iPaaS vs Traditional Integration", "https://www.sap.com/resources/ipaas-vs-traditional-migration"),
        ("Timus Consulting", "Direct API vs Middleware Integrations", "https://timusconsulting.com/direct-api-vs-middleware-integrations-which-approach-works-best-for-enterprise-systems/"),
        ("ERP Today", "SAP Shifts to AI Consumption Pricing", "https://erp.today/sap-shifts-to-ai-consumption-pricing-as-agents-threaten-saas-revenue-model/"),
        ("Workato", "SAP RFC Integration and Workflow Automation", "https://www.workato.com/integrations/sap"),
    ],
}

for category, refs in references.items():
    styled_heading(doc, category, level=2)
    for i, (source, title, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i}. ")
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r2 = p.add_run(f"[{source}] ")
        r2.bold = True
        r2.font.name = FONT_NAME
        r2.font.size = Pt(10)
        r3 = p.add_run(f'"{title}" ')
        r3.italic = True
        r3.font.name = FONT_NAME
        r3.font.size = Pt(10)
        r4 = p.add_run(url)
        r4.font.name = FONT_NAME
        r4.font.size = Pt(9)
        r4.font.color.rgb = ACCENT_COLOR

# ── Save ─────────────────────────────────────────────────────────────────────

output_path = "/home/shekerk/cc-best-practice-clone/automotive-ai-sap-odata-strategy-may2026.docx"
doc.save(output_path)

# Count references
total_refs = sum(len(refs) for refs in references.values())
print(f"Document saved to: {output_path}")
print(f"Total references: {total_refs}")
print("Done.")
